import logging
from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException, Header
from typing import Optional, List
from pydantic import BaseModel
import io

from app.services.rag_service import rag_service
from config import settings

logger = logging.getLogger(__name__)

router = APIRouter()

def verify_admin_key(x_admin_key: str = Header(None)):
    """Simple verification for admin endpoints"""
    # Dans un environnement réel, vérifier avec une vraie clé secrète configurée.
    # Ici, on utilise le SECRET_KEY de l'application ou un équivalent.
    if x_admin_key != settings.SECRET_KEY:
        raise HTTPException(status_code=403, detail="Non autorisé. Clé admin invalide.")
    return True

class TextDocumentRequest(BaseModel):
    title: str
    content: str
    source: str
    category: Optional[str] = None

@router.post("/documents/text", dependencies=[Depends(verify_admin_key)])
async def ingest_text_document(request: TextDocumentRequest):
    """
    Ingérer un document sous forme de texte brut.
    """
    metadata = {
        "title": request.title,
        "source": request.source,
        "category": request.category
    }
    
    result = await rag_service.ingest_document(content=request.content, metadata=metadata)
    if result["status"] == "error":
        raise HTTPException(status_code=500, detail=result["message"])
        
    return {"message": "Document indexé avec succès", "details": result}


@router.post("/documents/file", dependencies=[Depends(verify_admin_key)])
async def ingest_file_document(
    title: str = Form(...),
    source: str = Form(...),
    category: Optional[str] = Form(None),
    file: UploadFile = File(...)
):
    """
    Ingérer un document depuis un fichier (PDF, TXT, DOCX).
    """
    content = ""
    file_bytes = await file.read()
    
    try:
        if file.filename.endswith(".txt"):
            content = file_bytes.decode('utf-8')
            
        elif file.filename.endswith(".pdf"):
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
                
        elif file.filename.endswith(".docx"):
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
        else:
            raise HTTPException(status_code=400, detail="Format de fichier non supporté. Utilisez TXT, PDF ou DOCX.")
            
    except Exception as e:
        logger.error(f"Erreur d'extraction du fichier: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur d'extraction du fichier: {str(e)}")
        
    metadata = {
        "title": title,
        "source": source,
        "category": category,
        "filename": file.filename
    }
    
    result = await rag_service.ingest_document(content=content, metadata=metadata)
    if result["status"] == "error":
        raise HTTPException(status_code=500, detail=result["message"])
        
    return {"message": "Fichier indexé avec succès", "details": result}

@router.get("/documents", dependencies=[Depends(verify_admin_key)])
async def list_documents():
    """
    Lister les documents indexés.
    """
    documents = await rag_service.list_documents_summary()
    return {"documents": documents}

@router.delete("/documents/{source}", dependencies=[Depends(verify_admin_key)])
async def delete_document(source: str):
    """
    Supprimer tous les chunks associés à une source spécifique.
    """
    success = await rag_service.delete_document_by_source(source)
    if success:
        return {"message": f"Document(s) de la source '{source}' supprimé(s)."}
    return {"message": "Erreur lors de la suppression."}

@router.get("/documents/{source}/chunks", dependencies=[Depends(verify_admin_key)])
async def get_document_chunks(source: str):
    """
    Récupère tous les chunks (enregistrements) d'un document.
    """
    chunks = await rag_service.get_chunks_by_source(source)
    return {"chunks": chunks}

class UpdateChunkRequest(BaseModel):
    content: str

@router.put("/chunks/{chunk_id}", dependencies=[Depends(verify_admin_key)])
async def update_document_chunk(chunk_id: str, request: UpdateChunkRequest):
    """
    Met à jour le contenu d'un chunk et recalcule son embedding.
    """
    success = await rag_service.update_chunk(chunk_id, request.content)
    if success:
        return {"message": "Chunk mis à jour avec succès"}
    raise HTTPException(status_code=500, detail="Erreur lors de la mise à jour du chunk.")
